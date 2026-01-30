from typing import Dict, Any, Optional
from app.services.groq_service import get_groq_client, call_groq
from groq import GroqError
from app.services.usage_service import log_usage
from supabase import Client
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re # Import re for regex operations
import json
import logging

logger = logging.getLogger(__name__)


# Helper function to clean markdown text for docx (re-defined here for self-containment)
def _clean_markdown_text_for_docx(text_content: str) -> str:
    # Replace HTML <br> with newline
    text_content = text_content.replace('<br>', '\n')
    
    # Remove bold, italic, and strikethrough markers
    text_content = re.sub(r'(\*\*|__)(.*?)\1', r'\2', text_content) # **bold** or __bold__
    text_content = re.sub(r'(\*|_)(.*?)\1', r'\2', text_content)   # *italic* or _italic_
    text_content = re.sub(r'~~(.*?)~~', r'\1', text_content)       # ~~strikethrough~~

    # Remove links [text](url) -> text
    text_content = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text_content)

    # Remove inline code blocks `code`
    text_content = re.sub(r'`([^`]+)`', r'\1', text_content)

    # More aggressive cleanup for math environments for simpler display if not rendering
    text_content = re.sub(r'\$.*?\$', '', text_content) # Remove inline math $...$
    text_content = re.sub(r'\\[a-zA-Z]+', '', text_content) # Remove LaTeX commands like \frac, \sqrt
    text_content = re.sub(r'\{.*?\}', '', text_content) # Remove content in curly braces after LaTeX commands
    text_content = text_content.replace('$', '') # Catch any remaining lone $

    # Handle Markdown tables: simply strip pipes and header separators
    # This will turn tables into continuous lines of text, which is a compromise for simplicity
    text_content = re.sub(r'\|.*\|', lambda m: m.group(0).replace('|', ' '), text_content) # Replace pipes with spaces
    text_content = re.sub(r'[-=]+\s*[-=]+\s*[-=]+', '', text_content) # Remove table header separators (---)
    
    # Remove block code fences ```
    text_content = text_content.replace('```', '')

    return text_content.strip()


async def generate_course_outline(
    supabase: Client,
    user_id: str,
    username: str,
    course_full_name: str,
    course_code: Optional[str] = None,
    university_name: Optional[str] = None
) -> Dict[str, Any]:
    
    if not course_full_name:
        return {"success": False, "message": "Course Full Name is required."}

    client, error_message = get_groq_client()
    if error_message:
        return {"success": False, "message": error_message}
    
    uni_context = f"taught at {university_name}." if university_name else "taught at a major Nigerian University."
    code_context = f"(Code: {course_code})" if course_code else ""

    outline_prompt_content = f"""
    Generate a comprehensive, 12-week university_level course outline for the 
    course: "{course_full_name}" {code_context}. The course should reflect standards {uni_context}.

    **REQUIRED SECTIONS:**
    1. **Course Description:** (2-3 sentences)
    2. **Course Objectives:** (4-5 bullet points)
    3. **12-Week Schedule:** Use a **Markdown table** with the columns: **Week**, **Topic**, and **Key Learning Objectives**.

    Ensure the output is formatted cleanly using **Markdown**.
    """

    messages = [
        {"role": "system", "content": "You are an expert curriculum designer."},
        {"role": "user", "content": outline_prompt_content}
    ]

    try:
        response = None
        models = [
            "llama-3.1-8b-instant",
            "llama3-8b-8192"
        ]

        for model in models:
            try:
                response = call_groq(
                    client,
                    messages=messages,
                    model=model,
                    temperature=0.4
                )
                break
            except Exception as e:
                logger.warning(f"Groq model {model} failed for Course Outline: {e}")

        if not response:
            return {
                "success": False,
                "message": "AI service is currently overloaded. Please try again."
            }
        
        outline_text = response.choices[0].message.content.strip()

        await log_usage(
            supabase=supabase,
            user_id=user_id,
            user_name=username,
            feature_name="Course Outline Generator",
            action="generated",
            metadata={"course": course_full_name}
        )

        return {"success": True, "outline_text": outline_text}
    except GroqError as e:
        msg = str(e)
        if "429" in msg:
            return {"success": False, "message": "Too many requests. Please wait briefly."}
        logger.error(f"Groq API error during Course Outline generation: {msg}", exc_info=True)
        return {"success": False, "message": "AI service error. Please try again."}
    except Exception as e:
        logger.error("Unexpected error during course outline generation", exc_info=True)
        return {"success": False, "message": "An unexpected error occurred while generating the AI response."}

async def create_docx_from_outline(outline_text: str, course_full_name: str) -> io.BytesIO:
    doc = Document()
    doc.add_heading(f"Course Outline: {course_full_name}", 0) 
    
    for line in outline_text.split('\n'):
        stripped_line = line.strip()

        if not stripped_line: # Skip empty lines
            doc.add_paragraph("") # Add an empty paragraph for line breaks
            continue
        
        # Handle Headers (more robustly)
        header_match = re.match(r'^(#+)\s*(.*)', stripped_line)
        if header_match:
            level = len(header_match.group(1))
            text_content = header_match.group(2).strip()
            doc.add_heading(_clean_markdown_text_for_docx(text_content), level=min(level, 9)) # Max heading level in docx is 9
        # Handle Horizontal Rule
        elif re.match(r'^-{3,}$', stripped_line) or re.match(r'^\*{3,}$', stripped_line):
            doc.add_paragraph("-" * 20, style='Normal') # Add a simple line for HR
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Handle List Items
        elif re.match(r'^(\*|-|\+)\s', stripped_line):
            text_content = re.sub(r'^(\*|-|\+)\s', '', stripped_line).strip()
            doc.add_paragraph(_clean_markdown_text_for_docx(text_content), style='List Bullet')
        # Handle Blockquotes (simple paragraph with special formatting)
        elif stripped_line.startswith('>'):
            text_content = re.sub(r'^>\s*', '', stripped_line).strip()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(_clean_markdown_text_for_docx(text_content))
            run.italic = True # Simple blockquote style
        else:
            # All other content as normal paragraph
            text_content = _clean_markdown_text_for_docx(stripped_line)
            if text_content:
                doc.add_paragraph(text_content)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
