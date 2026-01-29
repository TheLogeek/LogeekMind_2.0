from typing import Dict, Any, List, Optional, Tuple
from app.services.groq_service import get_groq_client, call_groq
from groq import GroqError
from app.services.usage_service import log_usage, log_performance
from supabase import Client
from postgrest.exceptions import APIError
import json
from docx import Document
import io
import time
import re
import uuid
import datetime
import logging
from io import BytesIO
from PyPDF2 import PdfReader

logger = logging.getLogger(__name__)

# Configuration for chunking
MAX_CHUNK_SIZE = 5000  # Characters per chunk for lecture notes
CHUNK_OVERLAP = 400    # Overlap between chunks

# Helper function to extract text from file content
async def _extract_text_from_file_content(file_content: bytes, file_name: str) -> Optional[str]:
    """Extracts text from a file content based on its extension."""
    try:
        if file_name.lower().endswith('.pdf'):
            pdf_reader = PdfReader(BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            if not text.strip():
                return "Error: Could not extract text from PDF. The file might be image-based or corrupted."
            return text

        elif file_name.lower().endswith('.docx'):
            document = Document(BytesIO(file_content))
            text = ""
            for paragraph in document.paragraphs:
                text += paragraph.text + "\n"
            if not text.strip():
                return "Error: Could not extract text from DOCX. The file might be empty or corrupted."
            return text

        elif file_name.lower().endswith('.txt'):
            return file_content.decode("utf-8")
        else:
            return None
    except Exception as e:
        logger.error(f"Error extracting text from file {file_name}: {e}")
        return f"Error processing file {file_name}: {e}"


def create_intelligent_chunks(text: str, max_chunk_size: int = MAX_CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """
    Splits text into chunks intelligently, respecting paragraph and sentence boundaries.
    """
    if len(text) <= max_chunk_size:
        return [text]
    
    chunks = []
    paragraphs = text.split('\n\n')
    current_chunk = ""
    
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        
        if len(current_chunk) + len(paragraph) + 2 > max_chunk_size:
            if current_chunk:
                chunks.append(current_chunk.strip())
                overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
                current_chunk = overlap_text + "\n\n" + paragraph
            else:
                sentences = re.split(r'(?<=[.!?])\s+', paragraph)
                temp_chunk = ""
                
                for sentence in sentences:
                    if len(temp_chunk) + len(sentence) + 1 > max_chunk_size:
                        if temp_chunk:
                            chunks.append(temp_chunk.strip())
                            overlap_text = temp_chunk[-overlap:] if len(temp_chunk) > overlap else temp_chunk
                            temp_chunk = overlap_text + " " + sentence
                        else:
                            chunks.append(sentence[:max_chunk_size])
                            temp_chunk = sentence[max_chunk_size - overlap:]
                    else:
                        temp_chunk += " " + sentence if temp_chunk else sentence
                
                current_chunk = temp_chunk
        else:
            current_chunk += "\n\n" + paragraph if current_chunk else paragraph
    
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return chunks


async def summarize_lecture_notes_chunk(chunk: str, chunk_index: int, total_chunks: int, client: Any, model: str) -> Optional[str]:
    """
    Summarizes a chunk of lecture notes to extract key concepts.
    """
    context = f"This is part {chunk_index + 1} of {total_chunks} from lecture notes." if total_chunks > 1 else "These are complete lecture notes."
    
    system_prompt = "You are an expert at extracting key concepts from academic lecture notes."
    
    user_prompt = f"""{context}

Extract and summarize the key concepts, definitions, formulas, and important facts from this section.
Focus on information that would be suitable for exam questions.

Lecture notes section:
{chunk}

Provide a concise summary of the most important points."""

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]
    
    try:
        response = call_groq(
            client,
            messages=messages,
            model=model,
            temperature=0.2
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        logger.warning(f"Failed to summarize chunk {chunk_index + 1}: {e}")
        # Return first 800 chars as fallback
        return chunk[:800] + "..."


async def process_large_lecture_notes(lecture_notes_content: str, client: Any, model: str) -> str:
    """
    Processes large lecture notes by chunking and summarizing to fit model context.
    """
    text_length = len(lecture_notes_content)
    logger.info(f"Lecture notes length: {text_length} characters")
    
    # If notes are small enough, return as is
    if text_length <= MAX_CHUNK_SIZE:
        logger.info("Lecture notes fit within size limit")
        return lecture_notes_content
    
    # Large notes - need to chunk and summarize
    logger.info("Lecture notes exceed limit, applying chunking")
    chunks = create_intelligent_chunks(lecture_notes_content)
    logger.info(f"Created {len(chunks)} chunks from lecture notes")
    
    # Summarize each chunk to extract key concepts
    chunk_summaries = []
    for i, chunk in enumerate(chunks):
        logger.info(f"Processing chunk {i+1}/{len(chunks)}")
        summary = await summarize_lecture_notes_chunk(
            chunk=chunk,
            chunk_index=i,
            total_chunks=len(chunks),
            client=client,
            model=model
        )
        if summary:
            chunk_summaries.append(summary)
    
    # Combine summaries
    combined_summary = "\n\n---\n\n".join(
        [f"Section {i+1}:\n{summary}" for i, summary in enumerate(chunk_summaries)]
    )
    
    logger.info(f"Combined summary length: {len(combined_summary)} characters")
    return combined_summary


GRADE_POINTS = {
    "A": 5.0, "B": 4.0, "C": 3.0, "D": 2.0, "E": 1.0, "F": 0.0
}

def calculate_grade(score: int, total: int) -> Tuple[str, str, float]:
    if total == 0:
        return "N/A", "No questions graded.", 0.0
    
    percentage = (score / total) * 100
    if percentage >= 70:
        return "A", "Excellent! Distinction level.", percentage
    elif percentage >= 60:
        return "B", "Very Good. Keep it up.", percentage
    elif percentage >= 50:
        return "C", "Credit. You passed, but barely.", percentage
    elif percentage >= 45:
        return "D", "Pass. You need to study more.", percentage
    elif percentage >= 40:
        return "E", "Weak Pass. Dangerous territory.", percentage
    else:
        return "F", "Fail. You are not ready for this exam.", percentage

async def get_exam_performance_comparison(
    supabase: Client,
    shared_exam_id: str,
    current_score_percentage: float
) -> Dict[str, Any]:
    """
    Calculates how the current score compares to other submissions for the same exam.
    """
    try:
        response = supabase.table("shared_exam_submissions").select("percentage_score").eq("shared_exam_id", shared_exam_id).execute()
        
        all_percentages = [sub['percentage_score'] for sub in response.data if sub['percentage_score'] is not None]

        if not all_percentages:
            return {"success": True, "comparison_message": "No other submissions yet for comparison."}
        
        better_than_count = sum(1 for p in all_percentages if current_score_percentage > p)
        
        if len(all_percentages) > 0:
            percentile = (better_than_count / len(all_percentages)) * 100
            
            if percentile >= 90:
                comparison_message = f"Outstanding! You performed better than {percentile:.0f}% of test takers."
            elif percentile >= 75:
                comparison_message = f"Excellent! You performed better than {percentile:.0f}% of test takers."
            elif percentile >= 50:
                comparison_message = f"Good job! You performed better than {percentile:.0f}% of test takers."
            else:
                comparison_message = f"You performed better than {percentile:.0f}% of test takers. Keep studying!"

            return {"success": True, "comparison_message": comparison_message, "percentile": percentile}
        else:
            return {"success": True, "comparison_message": "Be the first to set the bar for this exam!"}

    except APIError as e:
        logger.error(f"Supabase APIError fetching exam submissions: {e.message}")
        return {"success": False, "message": "Could not retrieve comparison data."}
    except Exception as e:
        logger.error(f"Error calculating exam performance comparison: {e}", exc_info=True)
        return {"success": False, "message": "An unexpected error occurred during performance comparison."}


def validate_and_fix_exam_questions(exam_data: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Validates and fixes exam questions to ensure proper format.
    - Ensures 'answer' field contains ONLY the option letter (A, B, C, or D)
    - Fixes common malformations from AI
    - Handles both dictionary and object types for questions.
    """
    fixed_exam_data = []
    
    # Helper function to get value safely from dict or object
    def safe_get(item, key, default=None):
        if isinstance(item, dict):
            return item.get(key, default)
        else:
            return getattr(item, key, default)

    # Helper function to set value safely on dict or object
    def safe_set(item, key, value):
        if isinstance(item, dict):
            item[key] = value
        else:
            setattr(item, key, value)

    for q_idx, question in enumerate(exam_data):
        try:
            # Ensure required fields exist using safe_get
            # 'key in item' works for both dicts and objects if keys/attributes exist
            if not all(safe_get(question, key) is not None for key in ['question', 'options', 'answer', 'explanation']):
                logger.warning(f"Question {q_idx + 1} missing required fields, skipping")
                continue
            
            # Ensure options is a list with 4 items using safe_get
            options_val = safe_get(question, 'options')
            if not isinstance(options_val, list) or len(options_val) != 4:
                logger.warning(f"Question {q_idx + 1} has invalid options format, skipping")
                continue
            
            # Fix the answer field
            answer = safe_get(question, 'answer', '').strip()
            
            # If answer is already just a letter, keep it
            if len(answer) == 1 and answer.upper() in ['A', 'B', 'C', 'D']:
                safe_set(question, 'answer', answer.upper())
            else:
                # AI returned the full option text instead of letter - find which option matches
                answer_found = False
                for option_idx, option_text in enumerate(options_val):
                    if answer.lower() == option_text.lower() or answer in option_text or option_text in answer:
                        # Use safe_set for assignment
                        safe_set(question, 'answer', chr(65 + option_idx)) 
                        answer_found = True
                        # Log using safe_get to read the updated answer, handling both dict/object
                        logger.info(f"Fixed answer for Q{q_idx + 1}: '{answer}' -> '{safe_get(question, 'answer')}'")
                        break
                
                if not answer_found:
                    # Default to A if we can't determine the answer
                    logger.warning(f"Could not determine answer for Q{q_idx + 1}, defaulting to 'A'")
                    safe_set(question, 'answer', 'A')
            
            fixed_exam_data.append(question)
            
        except Exception as e:
            logger.error(f"Error validating question {q_idx + 1}: {e}")
            continue
    
    return fixed_exam_data


async def generate_exam_questions(
    supabase: Client,
    user_id: str,
    username: str,
    course_name: str,
    num_questions: int,
    topic: Optional[str] = None,
    lecture_notes_content: Optional[str] = None,
    file_name: Optional[str] = None,
    is_sharable: bool = False
) -> Dict[str, Any]:
    
    if not course_name:
        return {"success": False, "message": "Course Name is required."}
    
    if not topic and not lecture_notes_content:
        return {"success": False, "message": "Either a topic or lecture notes must be provided."}

    if is_sharable and user_id.startswith("guest_"):
        return {"success": False, "message": "Guest users cannot create sharable exams. Please log in to use this feature."}

    client, error_message = get_groq_client()
    if error_message:
        return {"success": False, "message": error_message}
    
    # Test which model is available
    models = ["llama-3.1-8b-instant", "mixtral-8x7b-32768"]
    working_model = None
    
    for model in models:
        try:
            test_response = call_groq(
                client,
                messages=[{"role": "user", "content": "Hi"}],
                model=model,
                temperature=0.1
            )
            working_model = model
            logger.info(f"Using model: {working_model}")
            break
        except Exception as e:
            logger.warning(f"Model {model} not available: {e}")
    
    if not working_model:
        return {"success": False, "message": "AI service is currently overloaded. Please try again."}
    
    # Process large lecture notes if provided
    if lecture_notes_content:
        lecture_notes_content = await process_large_lecture_notes(
            lecture_notes_content, 
            client, 
            working_model
        )
    
    # Construct prompt dynamically
    system_prompt = "You are an expert university professor setting an exam. You MUST follow the output format exactly."
    user_prompt_content = ""

    if lecture_notes_content:
        user_prompt_content = f"""
Generate {num_questions} examination-standard multiple-choice questions based ONLY on the provided lecture notes.

Course: {course_name}

Lecture Notes:
---
{lecture_notes_content}
---

CRITICAL INSTRUCTIONS:
1. Questions must be based ONLY on content from the lecture notes above
2. Each question must have exactly 4 options labeled A, B, C, D
3. The "answer" field MUST contain ONLY the letter (A, B, C, or D) - NOT the full text of the option
4. For mathematical questions, do NOT use LaTeX commands
5. Include at least one complex scenario or problem-solving question

OUTPUT FORMAT (STRICT):
Return ONLY a raw JSON array. Do NOT use markdown code blocks or any other formatting.
Each question must be a dictionary with these EXACT keys:
- "question": The question text (string)
- "options": Array of exactly 4 strings
- "answer": Single letter ONLY: "A", "B", "C", or "D" (string)
- "explanation": Why the answer is correct (string)

Example format:
[
  {{
    "question": "What is photosynthesis?",
    "options": ["Process of plant respiration", "Process converting light to energy", "Process of cell division", "Process of water absorption"],
    "answer": "B",
    "explanation": "Photosynthesis is the process by which plants convert light energy into chemical energy."
  }}
]

Now generate {num_questions} questions following this EXACT format:
"""
    else:
        user_prompt_content = f"""
Generate {num_questions} examination-standard multiple-choice questions.

Course: {course_name}
Topic: {topic if topic else 'General'}

CRITICAL INSTRUCTIONS:
1. Each question must have exactly 4 options labeled A, B, C, D
2. The "answer" field MUST contain ONLY the letter (A, B, C, or D) - NOT the full text of the option
3. For mathematical questions, do NOT use LaTeX commands
4. Questions should vary in difficulty
5. Include at least one complex scenario or problem-solving question

OUTPUT FORMAT (STRICT):
Return ONLY a raw JSON array. Do NOT use markdown code blocks or any other formatting.
Each question must be a dictionary with these EXACT keys:
- "question": The question text (string)
- "options": Array of exactly 4 strings
- "answer": Single letter ONLY: "A", "B", "C", or "D" (string)
- "explanation": Why the answer is correct (string)

Example format:
[
  {{
    "question": "What is the capital of France?",
    "options": ["London", "Paris", "Berlin", "Madrid"],
    "answer": "B",
    "explanation": "Paris is the capital and largest city of France."
  }}
]

Now generate {num_questions} questions following this EXACT format:
"""
    
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt_content}
    ]

    generated_exam_data = None
    share_id = None

    try:
        response = call_groq(
            client,
            messages=messages,
            model=working_model,
            temperature=0.4
        )
        
        response_content = response.choices[0].message.content.strip()
        
        if not response_content:
            logger.error("Groq API returned an empty response.")
            return {"success": False, "message": "AI returned an empty response. Please try again."}

        # Clean the response - remove markdown code blocks
        cleaned_text = response_content
        
        # Remove markdown code blocks
        cleaned_text = re.sub(r'```json\s*', '', cleaned_text)
        cleaned_text = re.sub(r'```\s*', '', cleaned_text)
        cleaned_text = cleaned_text.strip()
        
        # Try to extract JSON if there's extra text
        json_match = re.search(r'\[\s*\{.*\}\s*\]', cleaned_text, re.DOTALL)
        if json_match:
            cleaned_text = json_match.group(0)
        
        # Parse JSON
        generated_exam_data = json.loads(cleaned_text)
        
        # Validate it's a list
        if not isinstance(generated_exam_data, list):
            logger.error(f"Generated exam data is not a list: {type(generated_exam_data)}")
            return {"success": False, "message": "AI generated invalid exam format. Please try again."}
        
        # Validate and fix the exam questions
        generated_exam_data = validate_and_fix_exam_questions(generated_exam_data)
        
        if not generated_exam_data:
            logger.error("No valid questions after validation")
            return {"success": False, "message": "AI generated invalid questions. Please try again."}
        
        logger.info(f"Successfully generated {len(generated_exam_data)} valid questions")

        # Save to shared_exams if sharable
        if is_sharable:
            share_id = str(uuid.uuid4())
            try:
                supabase.table("shared_exams").insert({
                    "id": share_id,
                    "creator_id": user_id,
                    "title": f"{course_name} Exam ({len(generated_exam_data)} Qs)",
                    "exam_data": generated_exam_data
                }).execute()
                
            except APIError as db_e:
                logger.error(f"Supabase error saving shared exam: {db_e.message}")
                share_id = None
            except Exception as db_e:
                logger.error(f"Exception during Supabase insertion: {db_e}", exc_info=True)
                share_id = None

        await log_usage(
            supabase=supabase,
            user_id=user_id,
            user_name=username,
            feature_name="Exam Simulator",
            action="generated_exam",
            metadata={
                "course": course_name, 
                "topic": topic if topic else "notes", 
                "num_questions": len(generated_exam_data),
                "is_sharable": is_sharable, 
                "source_file": file_name
            }
        )

        return {"success": True, "exam_data": generated_exam_data, "share_id": share_id}

    except json.JSONDecodeError as e:
        logger.error(f"JSON Decode Error: {e}")
        logger.error(f"Response content: {response_content if response_content else 'No content'}")
        return {"success": False, "message": "AI generated an invalid exam format. Please try generating again."}
    except GroqError as e:
        msg = str(e)
        if "429" in msg:
            return {"success": False, "message": "Too many requests. Please wait briefly."}
        logger.error(f"Groq API error: {msg}", exc_info=True)
        return {"success": False, "message": "AI service error. Please try again."}
    except Exception as e:
        logger.error(f"Error during exam generation: {e}", exc_info=True)
        return {"success": False, "message": "An unexpected error occurred while generating the exam."}


async def grade_exam_and_log_performance(
    supabase: Client,
    user_id: str,
    username: str,
    exam_data: List[Any],
    user_answers: Dict[str, str],
    course_name: str,
    topic: Optional[str] = None,
    lecture_notes_source: bool = False
) -> Dict[str, Any]:

    score = 0
    total_questions = len(exam_data)

    for idx, q in enumerate(exam_data):
        user_selected_label = user_answers.get(str(idx))

        # ✅ SAFE extraction for dict OR object
        if isinstance(q, dict):
            correct_answer_letter = q.get("answer")
        else:
            correct_answer_letter = getattr(q, "answer", None)

        if not user_selected_label or not correct_answer_letter:
            continue

        user_selected_label = user_selected_label.strip().upper()
        correct_answer_letter = correct_answer_letter.strip().upper()

        if user_selected_label == correct_answer_letter:
            score += 1
            logger.debug(
                f"Q{idx}: Correct | User={user_selected_label} | Answer={correct_answer_letter}"
            )
        else:
            logger.debug(
                f"Q{idx}: Wrong | User={user_selected_label} | Answer={correct_answer_letter}"
            )

    grade, remark, percentage = calculate_grade(score, total_questions)

    await log_performance(
        supabase=supabase,
        user_id=user_id,
        feature="Exam Simulator",
        score=score,
        total_questions=total_questions,
        correct_answers=score,
        extra={
            "course": course_name,
            "topic": topic if topic else ("notes" if lecture_notes_source else "general"),
            "percentage": percentage
        }
    )

    await log_usage(
        supabase=supabase,
        user_id=user_id,
        user_name=username,
        feature_name="Exam Simulator",
        action="submitted_exam",
        metadata={
            "course": course_name,
            "score": score,
            "total": total_questions,
            "used_notes": lecture_notes_source,
            "percentage": percentage
        }
    )

    return {
        "success": True,
        "score": score,
        "total_questions": total_questions,
        "grade": grade,
        "remark": remark,
        "percentage": percentage
    }



async def create_docx_from_exam_results(
    exam_data: List[Dict[str, Any]],
    user_answers: Dict[str, str],
    score: int,
    total_questions: int,
    grade: str,
    course_name: str,
    topic: Optional[str] = None,
    lecture_notes_source: bool = False
) -> io.BytesIO:
    doc = Document()
    doc.add_heading(f"Exam Results: {course_name}", 0)
    
    if lecture_notes_source and topic:
        doc.add_paragraph(f"Source Context: {topic}")
    elif topic:
        doc.add_paragraph(f"Topic: {topic}")

    doc.add_paragraph(f"Final Score: {score}/{total_questions}\nGrade: {grade}")
    doc.add_paragraph("-" * 20)

    for idx, q in enumerate(exam_data):
        doc.add_heading(f"Question {idx + 1}", level=2)
        doc.add_paragraph(q['question'])
        
        doc.add_paragraph("Options:")
        for opt_idx, option in enumerate(q['options']):
            option_letter = chr(65 + opt_idx)
            doc.add_paragraph(f"  {option_letter}. {option}", style='List Bullet')
        
        user_answer_letter = user_answers.get(str(idx), "N/A").upper()
        correct_answer_letter = q.get('answer', 'N/A').upper()
        
        doc.add_paragraph(f"Your Answer: {user_answer_letter}")
        doc.add_paragraph(f"Correct Answer: {correct_answer_letter}")
        
        if user_answer_letter == correct_answer_letter:
            doc.add_paragraph("Result: ✓ Correct", style='Intense Quote')
        else:
            doc.add_paragraph("Result: ✗ Incorrect", style='Intense Quote')
        
        doc.add_paragraph(f"Explanation: {q.get('explanation', 'No explanation provided.')}")
        doc.add_paragraph("-" * 20)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


async def get_shared_exam(supabase: Client, share_id: str) -> Dict[str, Any]:
    """Fetches a shared exam and its creator's username."""
    try:
        response = supabase.table("shared_exams").select("*").eq("id", share_id).single().execute()
        
        creator_username = "A user"
        if response.data.get("creator_id"):
            try:
                profile_response = supabase.table("profiles").select("username").eq("id", response.data["creator_id"]).single().execute()
                if profile_response.data:
                    creator_username = profile_response.data.get("username", "A user")
            except APIError:
                pass
        
        return {
            "success": True, 
            "exam_data": response.data["exam_data"], 
            "creator_username": creator_username
        }
            
    except APIError as e:
        logger.error(f"Supabase APIError fetching shared exam {share_id}: {e.message}")
        return {"success": False, "message": "Exam not found or unavailable."}
    except Exception as e:
        logger.error(f"Error fetching shared exam {share_id}: {e}", exc_info=True)
        return {"success": False, "message": "A server error occurred while fetching the exam."}


async def get_shared_exam_submission_for_download(
    supabase: Client,
    user_id: str,
    shared_exam_id: str,
    submission_id: str
) -> Dict[str, Any]:
    """
    Fetches a specific shared exam submission for download.
    """
    try:
        submission_response = supabase.table("shared_exam_submissions").select("*").eq("id", submission_id).single().execute()
        
        if not submission_response.data:
            logger.warning(f"Submission {submission_id} not found.")
            return {"success": False, "message": "Submission not found."}
        
        submission = submission_response.data

        if submission.get("student_id") != user_id:
            logger.warning(f"Unauthorized download attempt for submission {submission_id}")
            return {"success": False, "message": "Unauthorized access to submission."}

        exam_fetch_response = await get_shared_exam(supabase, shared_exam_id)
        if not exam_fetch_response["success"]:
            return {"success": False, "message": exam_fetch_response.get("message", "Shared exam not found.")}
        
        shared_exam_title_response = supabase.table("shared_exams").select("title").eq("id", shared_exam_id).single().execute()
        if not shared_exam_title_response.data:
            logger.warning(f"Shared exam {shared_exam_id} title not found.")
            return {"success": False, "message": "Shared exam title not found."}

        exam_data = exam_fetch_response["exam_data"]
        course_name_and_topic = shared_exam_title_response.data.get("title", "Unknown Exam Topic")

        course_name_match = re.search(r"Course: (.*?)(?: - Topic:|$)", course_name_and_topic)
        topic_match = re.search(r"Topic: (.*?)(?: Exam|$)", course_name_and_topic)

        course_name = course_name_match.group(1).strip() if course_name_match else course_name_and_topic
        topic = topic_match.group(1).strip() if topic_match else None

        return {
            "success": True,
            "exam_data": exam_data,
            "score": submission["score"],
            "total_questions": submission["total_questions"],
            "grade": submission.get("grade", "N/A"),
            "course_name": course_name,
            "topic": topic,
            "user_answers": submission["user_answers"]
        }

    except APIError as e:
        logger.error(f"Supabase APIError fetching submission {submission_id}: {e.message}")
        return {"success": False, "message": "Failed to retrieve submission data."}
    except Exception as e:
        logger.error(f"Error fetching submission {submission_id}: {e}", exc_info=True)
        return {"success": False, "message": "An unexpected error occurred."}


async def submit_shared_exam_results(
    supabase: Client,
    share_id: str,
    user_answers: Dict[str, str],
    student_id: Optional[str] = None,
    student_identifier: Optional[str] = None
) -> Dict[str, Any]:
    """Grades and saves a submission for a shared exam."""
    try:
        exam_response = await get_shared_exam(supabase, share_id)
        if not exam_response["success"]:
            return exam_response

        exam_data = exam_response["exam_data"]
        total_questions = len(exam_data)
        
        # Grade the submission using fixed logic
        score = 0
        for idx, q in enumerate(exam_data):
            user_selected_label = user_answers.get(str(idx), "").strip().upper()
            correct_answer_letter = q['answer'].strip().upper()

            if user_selected_label == correct_answer_letter:
                score += 1
        
        grade, remark, percentage = calculate_grade(score, total_questions)

        submission_data = {
            "shared_exam_id": share_id,
            "student_id": student_id,
            "student_identifier": student_identifier,
            "user_answers": user_answers,
            "score": score,
            "total_questions": total_questions,
            "percentage_score": percentage,
            "grade": grade,
            "submitted_at": datetime.datetime.utcnow().isoformat() + "Z"
        }
        
        try:
            insert_response = supabase.table("shared_exam_submissions").insert(submission_data).execute()
            
            return {
                "success": True, 
                "submission_id": insert_response.data[0]['id'],
                "score": score,
                "total_questions": total_questions,
                "percentage_score": percentage,
                "grade": grade,
                "remark": remark
            }
        except APIError as db_e:
            logger.error(f"Supabase APIError submitting score: {db_e.message}")
            return {"success": False, "message": "Failed to save submission."}

    except Exception as e:
        logger.error(f"Error submitting shared exam score: {e}", exc_info=True)
        return {"success": False, "message": "A server error occurred while submitting your score."}
