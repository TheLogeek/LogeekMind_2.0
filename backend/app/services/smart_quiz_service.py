import json
import uuid
import io
import re
import logging
import datetime
from typing import Dict, Any, List, Optional, Tuple
from supabase import Client
from postgrest.exceptions import APIError
from app.services.groq_service import get_groq_client, call_groq
from groq import GroqError
from app.services.usage_service import log_usage, log_performance
from docx import Document

logger = logging.getLogger(__name__)

def _clean_markdown_text_for_docx(text_content: str) -> str:
    text_content = text_content.replace('<br>', '\n')
    text_content = re.sub(r'(\*\*|__)(.*?)\1', r'\2', text_content)
    text_content = re.sub(r'(\*|_)(.*?)\1', r'\2', text_content)
    text_content = re.sub(r'~~(.*?)~~', r'\1', text_content)
    text_content = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text_content)
    text_content = re.sub(r'`([^`]+)`', r'\1', text_content)
    text_content = re.sub(r'\[a-zA-Z]+\{.*?\}', '', text_content)
    text_content = re.sub(r'\[a-zA-Z]+', '', text_content)
    text_content = re.sub(r'\{.*?\}', '', text_content)
    text_content = text_content.replace('$', '')
    text_content = re.sub(r'\|.*\|', lambda m: m.group(0).replace('|', ' '), text_content)
    text_content = re.sub(r'[-=]+\s*[-=]+\s*[-=]+', '', text_content)
    text_content = text_content.replace('```', '')
    return text_content.strip()

DIFFICULTY_MAP = {1: "introductory", 2: "beginner", 3: "intermediate", 4: "advanced", 5: "expert"}

def validate_and_fix_quiz_questions(quiz_data: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    fixed_quiz_data = []
    for q_idx, question in enumerate(quiz_data):
        try:
            if not all(key in question for key in ['question', 'options', 'answer', 'explanation']):
                logger.warning(f"Question {q_idx + 1} missing required fields, skipping")
                continue
            
            if not isinstance(question['options'], list) or len(question['options']) not in [2, 4]:
                logger.warning(f"Question {q_idx + 1} has invalid options format, skipping")
                continue

            answer = question['answer'].strip()
            
            if len(answer) == 1 and answer.upper() in ['A', 'B', 'C', 'D']:
                question['answer'] = answer.upper()
            elif answer.lower() in ['true', 'false']:
                question['answer'] = answer.capitalize()
            else:
                answer_found = False
                for option_idx, option_text in enumerate(question['options']):
                    if answer.lower() == option_text.lower() or answer in option_text or option_text in answer:
                        if len(question['options']) == 4:
                            question['answer'] = chr(65 + option_idx)
                        else:
                            question['answer'] = option_text.capitalize()
                        answer_found = True
                        logger.info(f"Fixed answer for Q{q_idx + 1}: '{answer}' -> '{question['answer']}'")
                        break
                
                if not answer_found:
                    logger.warning(f"Could not determine answer for Q{q_idx + 1}, defaulting to 'A' or 'True'")
                    question['answer'] = 'A' if len(question['options']) == 4 else 'True'
            
            fixed_quiz_data.append(question)
            
        except Exception as e:
            logger.error(f"Error validating question {q_idx + 1}: {e}")
            continue
    return fixed_quiz_data

async def generate_quiz_service(
    supabase: Client,
    user_id: str,
    username: str,
    quiz_topic: str,
    num_questions: int,
    quiz_type: str,
    difficulty: int,
    is_sharable: bool = False
) -> Dict[str, Any]:

    if not quiz_topic:
        return {"success": False, "message": "Quiz topic is required."}

    if is_sharable and user_id.startswith("guest_"):
        return {"success": False, "message": "Guest users cannot create sharable quizzes. Please log in."}

    client, error_message = get_groq_client()
    if error_message:
        return {"success": False, "message": error_message}

    system_prompt = "You are an expert quiz creator. You MUST follow the output format exactly."
    
    options_instructions = """
- For "multiple-choice" questions, provide exactly 4 options labeled A, B, C, D. The "answer" MUST be the letter (e.g., "B").
- For "true-false" questions, the options MUST be ["True", "False"]. The "answer" MUST be "True" or "False".
"""

    quiz_prompt = f"""
Create a {quiz_type} quiz on the topic: "{quiz_topic}".
Difficulty: {DIFFICULTY_MAP[difficulty]}.
Number of Questions: {num_questions}.

CRITICAL INSTRUCTIONS:
{options_instructions}
- Each question must have a "question", "options", "answer", and "explanation".
- IMPORTANT: Never included a True/False question in a multiple choice quiz and never include a multiple choice question in a True/False quiz
- For mathematical questions, do NOT use LaTeX commands. Provide a detailed step-by-step solution in the "explanation".

OUTPUT FORMAT (STRICT):
Return ONLY a raw JSON array. Do NOT use markdown code blocks or any other formatting.
Each question must be a dictionary with these EXACT keys:
- "question": The question text (string)
- "options": An array of strings.
- "answer": The correct answer (string, either a letter or True/False).
- "explanation": A clear explanation of why the answer is correct (string).

Example for multiple-choice:
[
  {{
    "question": "What is the capital of France?",
    "options": ["London", "Paris", "Berlin", "Madrid"],
    "answer": "B",
    "explanation": "Paris is the capital and largest city of France."
  }}
]

Example for true-false:
[
  {{
    "question": "The Earth is flat.",
    "options": ["True", "False"],
    "answer": "False",
    "explanation": "The Earth is an oblate spheroid."
  }}
]

Now generate {num_questions} questions following this EXACT format:
"""

    try:
        response = None
        models = ["llama-3.1-8b-instant", "llama3-8b-8192"]

        for model in models:
            try:
                response = call_groq(
                    client,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": quiz_prompt}
                    ],
                    model=model,
                    temperature=0.4
                )
                break
            except Exception as e:
                logger.warning(f"Groq model {model} failed: {e}")

        if not response:
            return {"success": False, "message": "AI service is currently overloaded. Please try again."}

        content = response.choices[0].message.content.strip()
        
        cleaned_text = re.sub(r'```json\s*', '', content)
        cleaned_text = re.sub(r'```\s*', '', cleaned_text)
        cleaned_text = cleaned_text.strip()
        
        json_match = re.search(r'\[\s*\{.*\}\s*\]', cleaned_text, re.DOTALL)
        if json_match:
            cleaned_text = json_match.group(0)

        generated_quiz_data = json.loads(cleaned_text)

        if not isinstance(generated_quiz_data, list):
            logger.error(f"Generated quiz data is not a list: {type(generated_quiz_data)}")
            return {"success": False, "message": "AI generated invalid quiz format. Please try again."}

        generated_quiz_data = validate_and_fix_quiz_questions(generated_quiz_data)

        if not generated_quiz_data:
            logger.error("No valid questions after validation")
            return {"success": False, "message": "AI generated invalid questions. Please try again."}

        logger.info(f"Successfully generated {len(generated_quiz_data)} valid questions")

        share_id = None
        if is_sharable:
            share_id = str(uuid.uuid4())
            try:
                supabase.table("shared_quizzes").insert({
                    "id": share_id,
                    "creator_id": user_id,
                    "title": f"{quiz_topic} Quiz ({num_questions} Qs)",
                    "quiz_data": generated_quiz_data
                }).execute()
            except APIError as db_e:
                logger.error(f"Supabase APIError saving shared quiz: {db_e.message}")
                share_id = None
            except Exception as e:
                logger.error(f"Exception saving shared quiz: {e}", exc_info=True)
                share_id = None

        await log_usage(
            supabase=supabase,
            user_id=user_id,
            user_name=username,
            feature_name="Quiz Generator",
            action="generated",
            metadata={"topic": quiz_topic, "num_questions": len(generated_quiz_data), "is_sharable": is_sharable}
        )

        return {"success": True, "quiz_data": generated_quiz_data, "share_id": share_id}

    except GroqError as e:
        msg = str(e)
        if "429" in msg:
            return {"success": False, "message": "Too many requests. Please wait briefly."}
        logger.error(f"Groq API error during quiz generation: {msg}", exc_info=True)
        return {"success": False, "message": "AI service error. Please try again."}
    except json.JSONDecodeError:
        logger.error(f"Invalid JSON returned from Groq during quiz generation. Content: {content}", exc_info=True)
        return {"success": False, "message": "AI returned an invalid quiz format. Try again."}
    except Exception as e:
        logger.error("Unexpected error during quiz generation", exc_info=True)
        return {"success": False, "message": "An unexpected error occurred while generating the quiz."}

async def create_docx_from_quiz_results(
    quiz_data: List[Dict[str, Any]],
    quiz_topic: str,
    user_score: int,
    total_questions: int,
    user_answers: Dict[str, str]
) -> io.BytesIO:
    doc = Document()
    doc.add_heading(f"Quiz Results: {quiz_topic}", 0)
    doc.add_paragraph(f"Final Score: {user_score}/{total_questions}\n")

    for idx, q in enumerate(quiz_data):
        user_choice = user_answers.get(str(idx))
        question_text = q['question']
        doc.add_heading(f"Q{idx + 1}: {_clean_markdown_text_for_docx(question_text)}", level=2)

        doc.add_paragraph("Options:")
        for option in q['options']:
            doc.add_paragraph(_clean_markdown_text_for_docx(option), style='List Bullet')

        doc.add_paragraph(f"Your Answer: {_clean_markdown_text_for_docx(user_choice) if user_choice else '(No answer)'}")
        doc.add_paragraph(f"Correct Answer: {_clean_markdown_text_for_docx(q['answer'])}")
        doc.add_paragraph("Explanation:")
        explanation_text = q['explanation']
        for exp_line in explanation_text.split('\n'):
            stripped_exp_line = exp_line.strip()
            if stripped_exp_line:
                doc.add_paragraph(_clean_markdown_text_for_docx(stripped_exp_line))
        doc.add_paragraph("-" * 20)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

async def get_shared_quiz_submission_for_download(
    supabase: Client,
    user_id: str,
    shared_quiz_id: str,
    submission_id: str
) -> Dict[str, Any]:
    try:
        submission_response = await supabase.table("shared_quiz_submissions").select("*").eq("id", submission_id).single().execute()
        
        if not submission_response.data:
            logger.warning(f"Submission {submission_id} not found.")
            return {"success": False, "message": "Submission not found."}
        
        submission = submission_response.data

        if submission.get("student_id") != user_id:
            logger.warning(f"Unauthorized attempt to download submission {submission_id} by user {user_id}. Owner: {submission.get('student_id')}")
            return {"success": False, "message": "Unauthorized access to submission."}

        quiz_fetch_response = await get_shared_quiz(supabase, shared_quiz_id)
        if not quiz_fetch_response["success"]:
            return {"success": False, "message": quiz_fetch_response.get("message", "Shared quiz not found.")}
        
        shared_quiz_title_response = await supabase.table("shared_quizzes").select("title").eq("id", shared_quiz_id).single().execute()
        if not shared_quiz_title_response.data:
            logger.warning(f"Shared quiz {shared_quiz_id} title not found.")
            return {"success": False, "message": "Shared quiz title not found."}

        quiz_data = quiz_fetch_response["quiz_data"]
        quiz_topic = shared_quiz_title_response.data.get("title", "Unknown Quiz Topic")

        return {
            "success": True,
            "quiz_data": quiz_data,
            "quiz_topic": quiz_topic,
            "user_score": submission["score"],
            "total_questions": submission["total_questions"],
            "user_answers": submission["user_answers"]
        }

    except APIError as e:
        logger.error(f"Supabase APIError fetching submission {submission_id}: {e.message}")
        return {"success": False, "message": "Failed to retrieve submission data from database."}
    except Exception as e:
        logger.error(f"Error fetching shared quiz submission for download {submission_id}: {e}", exc_info=True)
        return {"success": False, "message": "An unexpected error occurred while preparing download."}

def calculate_grade(score: int, total: int) -> Tuple[str, str, float]:
    if total == 0:
        return "N/A", "No questions graded.", 0.0

    percentage = round((score / total) * 100, 2)
    if percentage >= 70:
        return "A", "Excellent! Distinction level.", percentage
    elif percentage >= 60:
        return "B", "Very Good. Keep it up.", percentage
    elif percentage >= 50:
        return "C", "Credit. You passed, but barely.", percentage
    elif percentage >= 45:
        return "D", "Pass. You need to study more.", percentage
    elif percentage >= 40:
        return "E", "Weak Pass. Dangerous territory.", percentage
    else:
        return "F", "Fail. You are not ready for this exam.", percentage

async def get_quiz_performance_comparison(
    supabase: Client,
    shared_quiz_id: str,
    current_score_percentage: float
) -> Dict[str, Any]:
    try:
        response = supabase.table("shared_quiz_submissions").select("percentage_score").eq("shared_quiz_id", shared_quiz_id).execute()
        
        all_percentages = [sub['percentage_score'] for sub in response.data if sub['percentage_score'] is not None]

        if not all_percentages:
            return {"success": True, "comparison_message": "No other submissions yet for comparison."}
        
        better_than_count = sum(1 for p in all_percentages if current_score_percentage > p)
        
        if len(all_percentages) > 0:
            percentile = (better_than_count / len(all_percentages)) * 100
            
            if percentile >= 90:
                comparison_message = f"Outstanding! You performed better than {percentile:.0f}% of test takers."
            elif percentile >= 75:
                comparison_message = f"Excellent! You performed better than {percentile:.0f}% of test takers."
            elif percentile >= 50:
                comparison_message = f"Good job! You performed better than {percentile:.0f}% of test takers."
            else:
                comparison_message = f"You performed better than {percentile:.0f}% of test takers. Keep studying!"

            return {"success": True, "comparison_message": comparison_message, "percentile": percentile}
        else:
            return {"success": True, "comparison_message": "Be the first to set the bar for this quiz!"}

    except APIError as e:
        logger.error(f"Supabase APIError fetching quiz submissions for comparison: {e.message}")
        return {"success": False, "message": "Could not retrieve comparison data."}
    except Exception as e:
        logger.error(f"Error calculating quiz performance comparison: {e}", exc_info=True)
        return {"success": False, "message": "An unexpected error occurred during performance comparison."}

async def get_shared_quiz(supabase: Client, share_id: str) -> Dict[str, Any]:
    """Fetches a shared quiz and its creator's username."""
    try:
        response = supabase.table("shared_quizzes").select("*").eq("id", share_id).single().execute()
        
        quiz_data = response.data
        creator_username = "A user"
        if quiz_data.get("creator_id"):
            try:
                profile_response = supabase.table("profiles").select("username").eq("id", quiz_data["creator_id"]).single().execute()
                if profile_response.data:
                    creator_username = profile_response.data.get("username", "A user")
            except APIError:
                pass
        
        quiz_data["creator_username"] = creator_username
        return {"success": True, **quiz_data}
            
    except APIError as e:
        logger.error(f"Supabase APIError fetching shared quiz {share_id}: {e.message}")
        return {"success": False, "message": "Quiz not found or unavailable."}
    except Exception as e:
        logger.error(f"Error fetching shared quiz {share_id}: {e}", exc_info=True)
        return {"success": False, "message": "A server error occurred while fetching the quiz."}

async def save_shared_quiz_submission(
    supabase: Client,
    shared_quiz_id: str,
    user_answers: Dict[str, str],
    student_id: Optional[str],
    student_identifier: Optional[str] = None
) -> Dict[str, Any]:
    try:
        quiz_fetch_response = await get_shared_quiz(supabase, shared_quiz_id)
        if not quiz_fetch_response["success"]:
            return quiz_fetch_response

        quiz_data = quiz_fetch_response["quiz_data"]
        total_questions = len(quiz_data)

        score = 0
        for idx, q in enumerate(quiz_data):
            user_selected_label = user_answers.get(str(idx), "").strip()
            correct_answer = q.get('answer', '').strip()

            if user_selected_label.lower() == correct_answer.lower():
                score += 1

        grade, remark, percentage = calculate_grade(score, total_questions)

        submission_data = {
            "shared_quiz_id": shared_quiz_id,
            "student_id": student_id,
            "student_identifier": student_identifier,
            "user_answers": user_answers,
            "score": score,
            "total_questions": total_questions,
            "percentage_score": percentage,
            "grade": grade,
            "submitted_at": datetime.datetime.utcnow().isoformat() + "Z"
        }
        
        try:
            response = supabase.table("shared_quiz_submissions").insert(submission_data).execute()

            return {
                "success": True, 
                "submission_id": response.data[0]['id'],
                "score": score,
                "total_questions": total_questions,
                "percentage_score": percentage,
                "grade": grade,
                "remark": remark
            }
        except APIError as db_e:
            logger.error(f"Supabase APIError submitting shared quiz score: {db_e.message}")
            return {"success": False, "message": "Failed to save submission."}

    except Exception as e:
        logger.error(f"Error submitting shared quiz score for quiz {shared_quiz_id}: {e}", exc_info=True)
        return {"success": False, "message": "A server error occurred while submitting your score."}
