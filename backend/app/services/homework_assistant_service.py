from typing import Dict, Any, Optional
from app.services.gemini_service import get_gemini_client
from google import genai
from app.services.usage_service import log_usage
from supabase import Client
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re # Import re for regex operations
from PIL import Image # Import PIL Image as in original utils.py


# Helper function to clean markdown text for docx
def _clean_markdown_text_for_docx(text_content: str) -> str:
    # Replace HTML <br> with newline
    text_content = text_content.replace('<br>', '\n')
    
    # Remove bold, italic, and strikethrough markers
    text_content = re.sub(r'(\*\*|__)(.*?)\1', r'\2', text_content) # **bold** or __bold__
    text_content = re.sub(r'(\*|_)(.*?)\1', r'\2', text_content)   # *italic* or _italic_
    text_content = re.sub(r'~~(.*?)~~', r'\1', text_content)       # ~~strikethrough~~

    # Remove links [text](url) -> text
    text_content = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text_content)

    # Remove inline code blocks `code`
    text_content = re.sub(r'`([^`]+)`', r'\1', text_content)

    # More aggressive cleanup for math environments for simpler display if not rendering
    text_content = re.sub(r'\$.*?\$', '', text_content) # Remove inline math $...$
    text_content = re.sub(r'\\[a-zA-Z]+', '', text_content) # Remove LaTeX commands like \frac, \sqrt
    text_content = re.sub(r'\{.*?\}', '', text_content) # Remove content in curly braces after LaTeX commands
    text_content = text_content.replace('$', '') # Catch any remaining lone $

    # Handle Markdown tables: simply strip pipes and header separators
    # This will turn tables into continuous lines of text, which is a compromise for simplicity
    text_content = re.sub(r'\|.*\|', lambda m: m.group(0).replace('|', ' '), text_content) # Replace pipes with spaces
    text_content = re.sub(r'[-=]+\s*[-=]+\s*[-=]+', '', text_content) # Remove table header separators (---)
    
    # Remove block code fences ```
    text_content = text_content.replace('```', '')

    return text_content.strip()


async def generate_homework_solution(
    supabase: Client,
    user_id: str,
    username: str,
    image_content: bytes, # Passed as bytes from FastAPI UploadFile
    image_mime_type: str, # mime_type is needed by PIL to identify the image format
    context: Optional[str] = None
) -> Dict[str, Any]:
    
    client, error_message = await get_gemini_client(user_id=user_id)
    if error_message:
        return {"success": False, "message": error_message}
    
    try:
        pil_image = Image.open(io.BytesIO(image_content))
    except Exception as e:
        return {"success": False, "message": f"Could not open image file: {e}"}
    
    full_prompt = f"""
    You are a rigorous academic solver. Based on the image and the user's instructions (if any),
    provide a complete and accurate solution. The output must be in a clean, professional, and easily 
    readable **Markdown format**.

    Instructions:
    {context if context else "The user provided no instructions"}
    """
    
    # Contents list is [PIL Image, prompt_text] as in the original Streamlit code
    contents = [pil_image, full_prompt]

    try:
        # Calling client.models.generate_content exactly as in the original utils.py and feature page
        response = client.models.generate_content(
            model="gemini-2.5-flash", # Using gemini-pro-vision for multimodal input
            contents=contents
        )
        
        solution_text = response.text

        await log_usage(
            supabase=supabase,
            user_id=user_id,
            user_name=username,
            feature_name="Homework Assistant",
            action="generated",
            metadata={"context": context}
        )

        return {"success": True, "solution_text": solution_text}

    except genai.errors.APIError as e:
        error_message = str(e)
        if "429" in error_message or "RESOURCE_EXHAUSTED" in error_message.upper():
            print(f"Gemini API rate limit exceeded during summarization: {e}")
            return "", "Gemini API rate limit exceeded. Please try again in a moment."
        elif "503" in error_message:
            print(f"AI is currently eperiencing high traffic. Try again shortly.")
            return "", "AI is currently eperiencing high traffic. Please try again shortly."
        else:
            print(f"An API error occurred: {e}")
            return "", f"An API error occurred: {e}"

    except Exception as e:
        print(f"Error during homework solution generation: {e}")
        return {"success": False, "message": "An unexpected error occurred while generating the solution."}

async def create_docx_from_solution(solution_text: str, context: Optional[str] = None) -> io.BytesIO:
    doc = Document()
    doc.add_heading("Homework Solution", 0)
    if context:
        doc.add_heading(f"Instructions: {context}", 1)
    
    for line in solution_text.split('\n'):
        stripped_line = line.strip()

        if not stripped_line: # Skip empty lines
            doc.add_paragraph("") # Add an empty paragraph for line breaks
            continue
        
        # Handle Headers (more robustly)
        header_match = re.match(r'^(#+)\s*(.*)', stripped_line)
        if header_match:
            level = len(header_match.group(1))
            text_content = header_match.group(2).strip()
            doc.add_heading(_clean_markdown_text_for_docx(text_content), level=min(level, 9)) # Max heading level in docx is 9
        # Handle Horizontal Rule
        elif re.match(r'^-{3,}$', stripped_line) or re.match(r'^\*{3,}$', stripped_line):
            doc.add_paragraph("-" * 20, style='Normal') # Add a simple line for HR
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Handle List Items
        elif re.match(r'^(\*|-|\+)\s', stripped_line):
            text_content = re.sub(r'^(\*|-|\+)\s', '', stripped_line).strip()
            doc.add_paragraph(_clean_markdown_text_for_docx(text_content), style='List Bullet')
        # Handle Blockquotes (simple paragraph with special formatting)
        elif stripped_line.startswith('>'):
            text_content = re.sub(r'^>\s*', '', stripped_line).strip()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(_clean_markdown_text_for_docx(text_content))
            run.italic = True # Simple blockquote style
        else:
            # All other content as normal paragraph
            text_content = _clean_markdown_text_for_docx(stripped_line)
            if text_content:
                doc.add_paragraph(text_content)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
